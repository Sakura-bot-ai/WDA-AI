import requests
from docx import Document
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import qn, nsmap
from configparser import ConfigParser
import os
import time
from tkinter import *
from tkinter import ttk, messagebox, colorchooser
import threading
from formatting import DocumentFormatter
from line import LineFormatter  # 新增导入
from word import ContentFilter  # 新增导入
from tkinter import filedialog

class AIDocWriter:
    def __init__(self):
        self.config = ConfigParser()
        self.config.read('config.ini')
        
        self.api_key = self.config.get('API', 'api_key')
        self.model = self.config.get('API', 'model_name')
        self.api_base = self.config.get('API', 'api_base')
        self.font_name = self.config.get('Style', 'font_name', fallback='微软雅黑')
        self.font_size = self.config.getint('Style', 'font_size', fallback=12)

    def create_document(self, content, filename, font_name=None, font_size=None, bold=False, italic=False, template=None, spacing_enabled=False):
        # 文件名冲突检测
        base_name, ext = os.path.splitext(filename)
        counter = 1
        while os.path.exists(filename):
            filename = f"{base_name}_{counter}{ext}"
            counter += 1
        if template and os.path.exists(template):
            doc = Document(template)
        else:
            doc = Document()
        
        # 使用主文档的格式化器实例
        formatter = DocumentFormatter(
            font_name=font_name or self.font_name,
            font_size=DocumentFormatter.configure_font(font_size) if font_size else self.font_size
        )
        
        try:
            # 清理段落装饰线时保留边框
            line_formatter = LineFormatter(spacing_enabled=True)
            for para in doc.paragraphs:
                formatter.clean_decorative_lines(para._element.get_or_add_pPr(), keep_borders=True)
                # 符号触发边框
                if ':' in para.text:
                    line_formatter.add_symbol_triggered_border(para, ':')
                # 段落分割线
                line_formatter.set_spacing_property(para)
            
            # 清理表格装饰线时保留边框
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        formatter.clean_decorative_lines(cell._element.get_or_add_tcPr(), keep_borders=True)
            
            # 强制设置分节符类型为连续
            for section in doc.sections:
                section.start_type = 4  # 直接设置CONTINUOUS类型
                sectPr = section._sectPr
                type_elements = sectPr.xpath('./w:type')
                if type_elements:
                    type_element = type_elements[0]
                    type_element.set(qn('w:val'), 'continuous')
                else:
                    type_element = parse_xml(f'<w:type xmlns:w="{nsmap["w"]}" w:val="continuous"/>')
                    sectPr.append(type_element)
            
            # 清理残留的XML装饰元素
            for element in doc.element.xpath('//*[contains(name(), "Bdr")]'):
                element.getparent().remove(element)
            
        except Exception as e:
            print(f"文档处理异常: {type(e).__name__} - {str(e)}\n发生位置: {e.__traceback__.tb_lineno if e.__traceback__ else '未知'}")
            raise RuntimeError("文档生成失败，请检查模板格式") from e
        
        # 在文档末尾添加新内容
        line_formatter = LineFormatter(spacing_enabled=spacing_enabled)
        for para_text in content.split('\n'):
            p = formatter.create_paragraph(doc, para_text, bold=bold, italic=italic)
            # 添加间隔线
            line_formatter.set_spacing_border(p)
            line_formatter.set_spacing_property(p)
        
        # 统一清理所有分节符
        for section in doc.sections:
            section.start_type = 4  # 4对应CONTINUOUS类型
        
        doc.save(filename)
        print(f"文档已保存至：{os.path.abspath(filename)}")
        return filename

    def generate_content(self, prompt):
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }

        payload = {
            "model": self.model,
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.7
        }

        try:
            response = requests.post(
                f"{self.api_base}/chat/completions",
                headers=headers,
                json=payload
            )
            response.raise_for_status()
            filtered_content = ContentFilter.filter_ai_symbols(response.json()['choices'][0]['message']['content'])
            return filtered_content
        except Exception as e:
            print(f"API请求失败: {e}")
            return None

class AIDocApp:
    def __init__(self):
        self.writer = AIDocWriter()
        
        self.window = Tk()
        self.window.title('AI文档生成器')
        self.window.geometry('500x300')
        
        self.create_widgets()

    def create_widgets(self):
        # 输入标签
        Label(self.window, text='请输入文档主题:').pack(pady=10)
        
        # 文本输入框
        self.prompt_input = Text(self.window, height=8, width=50)
        self.prompt_input.pack(pady=5)
        
        # 字体设置面板
        font_frame = Frame(self.window)
        font_frame.pack(pady=5)

        Label(font_frame, text='字体:').grid(row=0, column=0)
        # 原代码中toggle_theme方法定义在文件末尾，在create_widgets方法中调用时还未定义，将其移到类开头解决此问题
        def toggle_theme(self):
            current_bg = self.window.cget('bg')
            is_dark = current_bg == 'SystemButtonFace'
            
            theme_colors = {
                'dark': {
                    'bg': '#2d2d2d',
                    'fg': '#ffffff',
                    'entry_bg': '#3d3d3d',
                    'button_bg': '#4d4d4d'
                },
                'light': {
                    'bg': 'SystemButtonFace',
                    'fg': 'SystemWindowText',
                    'entry_bg': 'SystemWindow',
                    'button_bg': 'SystemButtonFace'
                }
            }
            
            colors = theme_colors['dark'] if is_dark else theme_colors['light']
            
            self.window.config(bg=colors['bg'])
            
            # 统一更新所有组件样式
            for widget in self.window.winfo_children():
                if isinstance(widget, (Text, Entry)):
                    widget.config(bg=colors['entry_bg'], fg=colors['fg'], insertbackground=colors['fg'])
                elif isinstance(widget, Button):
                    widget.config(bg=colors['button_bg'], fg=colors['fg'])
                elif isinstance(widget, Label):
                    widget.config(bg=colors['bg'], fg=colors['fg'])

            # 更新ttk组件样式（Windows专属）
            if os.name == 'nt':
                self.style = ttk.Style()
                self.style.theme_use('vista')
                self.style.configure('.', 
                    background=colors['bg'],
                    foreground=colors['fg'],
                    fieldbackground=colors['entry_bg']
                )
                self.style.map('TButton',
                    background=[('active', colors['button_bg'])],
                    foreground=[('active', colors['fg'])]
                )
                self.style.configure('TFrame', background=colors['bg'])
                self.style.configure('TButton', 
                    background='#404040' if is_dark else 'SystemButtonFace',
                    foreground=colors['fg']
                )
                self.style.configure('TCombobox', 
                    fieldbackground=colors['entry_bg'],
                    background=colors['entry_bg'],
                    foreground=colors['fg']
                )
        AIDocApp.toggle_theme = toggle_theme
        ttk.Button(font_frame, text='切换主题', command=lambda: self.toggle_theme()).grid(row=0, column=0, padx=2)
        self.font_combobox = ttk.Combobox(font_frame, values=['微软雅黑', '宋体', '黑体', '楷体', '仿宋', '思源黑体', '华文细黑', '方正姚体', 'Times New Roman', 'Arial', 'Georgia', 'Calibri'], width=12)
        self.font_combobox.set('微软雅黑')
        self.font_combobox.grid(row=0, column=1)

        Label(font_frame, text='字号:').grid(row=0, column=2, padx=5)
        self.font_size_combo = ttk.Combobox(font_frame, values=['八号(5)', '七号(5.5)', '小六(6.5)', '六号(7.5)', '小五(9)', '五号(10.5)', '小四(12)', '四号(14)', '小三(15)', '三号(16)', '小二(18)', '二号(22)', '小一(24)', '一号(26)', '小初(36)', '初号(42)'], width=10)
        self.font_size_combo.set('小四(12)')
        self.font_size_combo.grid(row=0, column=3)

        self.bold_var = BooleanVar()
        self.italic_var = BooleanVar()
        self.spacing_var = BooleanVar()
        self.symbol_filter_var = BooleanVar()  # 新增间隔线状态变量
        ttk.Checkbutton(font_frame, text='B', command=self.toggle_bold, variable=self.bold_var).grid(row=0, column=4, padx=2)
        ttk.Checkbutton(font_frame, text='I', command=self.toggle_italic, variable=self.italic_var).grid(row=0, column=5, padx=2)
        ttk.Checkbutton(font_frame, text='间隔线', variable=self.spacing_var).grid(row=0, column=6, padx=2)
        ttk.Checkbutton(font_frame, text='符号过滤', variable=self.symbol_filter_var).grid(row=0, column=7, padx=2)

        # 模板选择面板
        template_frame = Frame(self.window)
        template_frame.pack(pady=5)
        
        Label(template_frame, text='文档模板:').grid(row=0, column=0)
        self.template_entry = Entry(template_frame, width=35)
        self.template_entry.grid(row=0, column=1, padx=5)
        
        # 初始化文档生成参数
        self.doc_params = {
            'font_name': self.font_combobox.get(),
            'font_size': self.font_size_combo.get(),
            'bold': self.bold_var.get(),
            'italic': self.italic_var.get(),
            'spacing_enabled': self.spacing_var.get(),
            'template': self.template_entry.get()
        }
        
        ttk.Button(template_frame, text='浏览模板', command=self.select_template).grid(row=0, column=2)

        # 初始化文档生成参数
        self.doc_params = {
            'font_name': self.font_combobox.get(),
            'font_size': self.font_size_combo.get(),
            'bold': self.bold_var.get(),
            'italic': self.italic_var.get(),
            'spacing_enabled': self.spacing_var.get(),
            'template': self.template_entry.get()
        }

        # 模板预览区域
        self.preview_text = Text(self.window, height=6, width=50)
        self.preview_text.pack(pady=5)
        self.preview_text.config(state=DISABLED)

        ttk.Button(template_frame, text='浏览模板', command=self.select_template).grid(row=0, column=2)

        # 生成按钮
        self.generate_button = Button(self.window, text='生成文档', command=self.generate_document)
        self.generate_button.pack(pady=10, side=BOTTOM)

        # 状态标签
        self.status_label = Label(self.window, text='', fg='black')
        self.status_label.pack(pady=5)

    def toggle_bold(self):
        current_font = self.font_combobox.get()
        self.font_combobox.config(font=(current_font, 10, 'bold' if self.bold_var.get() else ''))

    def toggle_italic(self):
        current_font = self.font_combobox.get()
        self.font_combobox.config(font=(current_font, 10, 'italic' if self.italic_var.get() else ''))

    def generate_document(self):
        self.generate_button.config(state=DISABLED)
        self.status_label.config(text='生成中...', fg='blue')
        
        prompt = self.prompt_input.get('1.0', END).strip()
        font_params = {
            'font_name': self.font_combobox.get(),
            'font_size': self.font_size_combo.get(),
            'bold': self.bold_var.get(),
            'italic': self.italic_var.get(),
            'spacing_enabled': self.spacing_var.get(),  # 新增间隔线参数
            'template': self.template_entry.get()
        }

        def api_thread():
            try:
                generated_content = self.writer.generate_content(self.prompt_input.get('1.0', END).strip())
                if generated_content:
                    filename = self.writer.create_document(generated_content, 'AI生成的文档.docx', **font_params)
                    self.window.after(0, lambda: self.status_label.config(
                        text=f'文档已生成: {os.path.abspath(filename)}',
                        fg='green'
                    ))
                else:
                    self.window.after(0, lambda: self.status_label.config(
                        text='生成失败：API返回空内容',
                        fg='red'
                    ))
            except Exception as e:
                self.window.after(0, lambda e=e: self.status_label.config(
                    text=f'生成失败：{str(e)}',
                    fg='red'
                ))
            finally:
                self.window.after(0, lambda: self.generate_button.config(state=NORMAL))

        threading.Thread(target=api_thread, daemon=True).start()

    
    def select_template(self):
        # 确保示例目录存在
        os.makedirs('data/example', exist_ok=True)
        filepath = filedialog.askopenfilename(
            initialdir='data/example',
            title='选择示例文档',
            filetypes=[('Word文档', '*.docx')]
        )
        if filepath:
            self.template_entry.delete(0, END)
            self.template_entry.insert(0, filepath)
            if os.path.exists(filepath):
                self.update_preview(filepath)
            else:
                messagebox.showerror('错误', '模板文件不存在或路径错误')

    def update_preview(self, filepath):
        try:
            if not os.path.exists(filepath):
                raise FileNotFoundError(f'文件 {filepath} 不存在')
            doc = Document(filepath)
            preview_content = '\n'.join([p.text for p in doc.paragraphs[:10]])
            self.preview_text.config(state=NORMAL)
            self.preview_text.delete(1.0, END)
            self.preview_text.insert(END, preview_content + '\n...（预览截取前10段内容）')
            self.preview_text.config(state=DISABLED)
        except Exception as e:
            messagebox.showerror('预览错误', f'无法读取模板文件: {str(e)}')

if __name__ == "__main__":
    app = AIDocApp()
    app.window.mainloop()

    def toggle_theme(self):
        current_bg = self.window.cget('bg')
        is_dark = current_bg == 'SystemButtonFace'
        
        theme_colors = {
            'dark': {
                'bg': '#2d2d2d',
                'fg': '#ffffff',
                'entry_bg': '#3d3d3d',
                'button_bg': '#4d4d4d'
            },
            'light': {
                'bg': 'SystemButtonFace',
                'fg': 'SystemWindowText',
                'entry_bg': 'SystemWindow',
                'button_bg': 'SystemButtonFace'
            }
        }
        
        colors = theme_colors['dark'] if is_dark else theme_colors['light']
        
        self.window.config(bg=colors['bg'])
        
        # 统一更新所有组件样式
        for widget in self.window.winfo_children():
            if isinstance(widget, (Text, Entry)):
                widget.config(bg=colors['entry_bg'], fg=colors['fg'], insertbackground=colors['fg'])
            elif isinstance(widget, Button):
                widget.config(bg=colors['button_bg'], fg=colors['fg'])
            elif isinstance(widget, Label):
                widget.config(bg=colors['bg'], fg=colors['fg'])

        # 更新ttk组件样式（Windows专属）
        if os.name == 'nt':
            self.style = ttk.Style()
            self.style.theme_use('vista')
            self.style.configure('.', 
                background=colors['bg'],
                foreground=colors['fg'],
                fieldbackground=colors['entry_bg']
            )
            self.style.map('TButton',
                background=[('active', colors['button_bg'])],
                foreground=[('active', colors['fg'])]
            )
            self.style.configure('TFrame', background=colors['bg'])
            self.style.configure('TButton', 
                background='#404040' if is_dark else 'SystemButtonFace',
                foreground=colors['fg']
            )
            self.style.configure('TCombobox', 
                fieldbackground=colors['entry_bg'],
                background=colors['entry_bg'],
                foreground=colors['fg']
            )