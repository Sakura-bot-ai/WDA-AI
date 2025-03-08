from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml.ns import nsmap

import copy

class DocumentFormatter:
    def __init__(self, font_name='微软雅黑', font_size=12, template_doc=None):
        self.font_name = font_name
        self.font_size = Pt(font_size)
        self.template_doc = template_doc

    def clean_decorative_lines(self, element, keep_borders=False):
        # 清理装饰线时保留边框
        if not keep_borders:
            for border_attr in ['top', 'left', 'bottom', 'right', 'between']:
                elem = element.find(qn(f'w:{border_attr}'))
                if elem is not None:
                    element.remove(elem)
        
        # 清理段落边框
        pBdr = element.find(qn('w:pBdr'))
        if pBdr is not None:
            element.remove(pBdr)

    def merge_style_properties(self, target_style, source_style):
        # 跳过基础样式的覆盖并清理边框属性
        if source_style.name == 'Normal' and target_style.name == 'Normal':
            return
        
        # 清理源样式的边框属性
        self.clean_decorative_lines(source_style.element)
        
        # 深度合并字体属性
        source_font = source_style.font
        target_font = target_style.font
        
        for attr in ['name', 'size', 'color', 'bold', 'italic', 'underline']:
            value = getattr(source_font, attr, None)
            if value not in [None, False]:
                setattr(target_font, attr, value)
        
        # 合并段落格式
        source_para = source_style.paragraph_format
        target_para = target_style.paragraph_format
        for attr in ['alignment', 'first_line_indent', 'line_spacing', 'space_before', 'space_after']:
            value = getattr(source_para, attr, None)
            if value is not None:
                setattr(target_para, attr, value)
        # 清理段落边框属性
        if hasattr(source_para, 'borders'):
            target_para.border = None
        if source_para.element.pBdr is not None:
            target_para.element.pPr.remove_all('pBdr')

    def copy_paragraph_styles(self, target_doc, template_doc, paragraph, **kwargs):
        if template_doc:
            style_name = paragraph.style.name
            # 特殊处理基础样式
            if style_name in ['Normal', 'Heading 1', 'Heading 2']:
                self.merge_style_properties(target_doc.styles[style_name], template_doc.styles[style_name])
                return
                
            template_style = template_doc.styles[style_name]
            if template_style.name not in target_doc.styles:
                new_style = target_doc.styles.add_style(template_style.name, template_style.type)
                new_style.element.append(copy.deepcopy(template_style.element))
            else:
                self.merge_style_properties(target_doc.styles[template_style.name], template_style)
            if kwargs.get('bold'):
                new_style.font.bold = True
            if kwargs.get('italic'):
                new_style.font.italic = True

    def apply_base_styles(self, doc):
        if self.template_doc:
            for style in self.template_doc.styles:
                # 检查样式是否已存在
                if style.name not in doc.styles:
                    # 深拷贝样式元素
                    new_style = doc.styles.add_style(style.name, style.type)
                    new_style.element.append(copy.deepcopy(style.element))
                else:
                    # 合并样式属性
                    existing_style = doc.styles[style.name]
                    self.merge_style_properties(existing_style, style)
        else:
            # 保留原有基础样式设置
            style = doc.styles['Normal']
            font = style.font
            font.name = self.font_name
            font.size = self.font_size
            # 添加中文字体支持
            font.element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)

    def create_paragraph(self, doc, text, bold=False, italic=False, style_name=None):
        if style_name and style_name in doc.styles:
            p = doc.add_paragraph(style=style_name)
        else:
            p = doc.add_paragraph()
        
        run = p.add_run(text)
        run.font.name = self.font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
        run.bold = bold
        run.italic = italic
        return p

    @staticmethod
    def configure_font(font_size_str):
        return Pt(float(font_size_str.split('(')[1].replace(')','')))